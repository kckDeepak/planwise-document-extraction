"""
Script to count documents and pages per client/folder.
Supports PDF and Word documents (.docx, .doc)
"""

import os
from pathlib import Path
from collections import defaultdict
import json
from datetime import datetime

# Try to import PDF library
try:
    from pypdf import PdfReader
    PDF_SUPPORT = True
except ImportError:
    try:
        from PyPDF2 import PdfReader
        PDF_SUPPORT = True
    except ImportError:
        PDF_SUPPORT = False
        print("Warning: pypdf or PyPDF2 not installed. PDF page counting will be unavailable.")
        print("Install with: pip install pypdf")

# Try to import docx library
try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    print("Warning: python-docx not installed. DOCX page counting will be unavailable.")
    print("Install with: pip install python-docx")


def count_pdf_pages(file_path: str) -> int:
    """Count the number of pages in a PDF file."""
    if not PDF_SUPPORT:
        return 0
    try:
        reader = PdfReader(file_path)
        return len(reader.pages)
    except Exception as e:
        print(f"  Warning: Could not read PDF '{file_path}': {e}")
        return 0


def count_docx_pages(file_path: str) -> int:
    """
    Estimate the number of pages in a DOCX file.
    Note: DOCX files don't have a fixed page count - it depends on rendering.
    This estimates based on paragraph breaks and content.
    """
    if not DOCX_SUPPORT:
        return 0
    try:
        doc = Document(file_path)
        # Count total paragraphs and estimate pages
        # Average ~25-30 paragraphs per page is a rough estimate
        total_paragraphs = len(doc.paragraphs)
        # More accurate: count characters and estimate (~3000 chars per page)
        total_chars = sum(len(p.text) for p in doc.paragraphs)
        estimated_pages = max(1, total_chars // 3000) if total_chars > 0 else 1
        return estimated_pages
    except Exception as e:
        print(f"  Warning: Could not read DOCX '{file_path}': {e}")
        return 0


def analyze_folder(root_path: str) -> dict:
    """
    Analyze the folder structure and count documents/pages per client folder.
    
    Returns a dictionary with client folder stats.
    """
    root_path = Path(root_path)
    
    if not root_path.exists():
        raise FileNotFoundError(f"Path does not exist: {root_path}")
    
    results = {}
    total_stats = {
        "total_pdf_files": 0,
        "total_docx_files": 0,
        "total_pdf_pages": 0,
        "total_docx_pages": 0,
        "total_documents": 0,
        "total_pages": 0
    }
    
    # Get all client folders (immediate subdirectories)
    client_folders = [f for f in root_path.iterdir() if f.is_dir()]
    
    print(f"\nAnalyzing {len(client_folders)} client folders in: {root_path}\n")
    print("=" * 80)
    
    for client_folder in sorted(client_folders):
        client_name = client_folder.name
        print(f"\nProcessing: {client_name}")
        
        client_stats = {
            "pdf_files": [],
            "docx_files": [],
            "pdf_count": 0,
            "docx_count": 0,
            "pdf_pages": 0,
            "docx_pages": 0,
            "total_documents": 0,
            "total_pages": 0
        }
        
        # Walk through all subdirectories of this client folder
        for root, dirs, files in os.walk(client_folder):
            for file in files:
                file_path = Path(root) / file
                relative_path = file_path.relative_to(client_folder)
                
                # Check for PDF files
                if file.lower().endswith('.pdf'):
                    pages = count_pdf_pages(str(file_path))
                    client_stats["pdf_files"].append({
                        "file": str(relative_path),
                        "pages": pages
                    })
                    client_stats["pdf_count"] += 1
                    client_stats["pdf_pages"] += pages
                
                # Check for Word documents
                elif file.lower().endswith(('.docx', '.doc')):
                    if file.lower().endswith('.docx'):
                        pages = count_docx_pages(str(file_path))
                    else:
                        # .doc files can't be read with python-docx
                        pages = 0
                        print(f"  Note: .doc file skipped for page count: {file}")
                    
                    client_stats["docx_files"].append({
                        "file": str(relative_path),
                        "pages": pages
                    })
                    client_stats["docx_count"] += 1
                    client_stats["docx_pages"] += pages
        
        # Calculate totals for this client
        client_stats["total_documents"] = client_stats["pdf_count"] + client_stats["docx_count"]
        client_stats["total_pages"] = client_stats["pdf_pages"] + client_stats["docx_pages"]
        
        results[client_name] = client_stats
        
        # Update global totals
        total_stats["total_pdf_files"] += client_stats["pdf_count"]
        total_stats["total_docx_files"] += client_stats["docx_count"]
        total_stats["total_pdf_pages"] += client_stats["pdf_pages"]
        total_stats["total_docx_pages"] += client_stats["docx_pages"]
        total_stats["total_documents"] += client_stats["total_documents"]
        total_stats["total_pages"] += client_stats["total_pages"]
        
        print(f"  PDFs: {client_stats['pdf_count']} ({client_stats['pdf_pages']} pages)")
        print(f"  Word: {client_stats['docx_count']} ({client_stats['docx_pages']} pages estimated)")
        print(f"  Total: {client_stats['total_documents']} documents, {client_stats['total_pages']} pages")
    
    return {
        "clients": results,
        "summary": total_stats,
        "analyzed_path": str(root_path),
        "analyzed_at": datetime.now().isoformat()
    }


def print_summary(results: dict):
    """Print a formatted summary of the analysis."""
    print("\n" + "=" * 80)
    print("SUMMARY")
    print("=" * 80)
    
    summary = results["summary"]
    clients = results["clients"]
    
    print(f"\nTotal Client Folders: {len(clients)}")
    print(f"\nDocument Counts:")
    print(f"  PDF Files:  {summary['total_pdf_files']}")
    print(f"  Word Files: {summary['total_docx_files']}")
    print(f"  Total:      {summary['total_documents']}")
    
    print(f"\nPage Counts:")
    print(f"  PDF Pages:  {summary['total_pdf_pages']}")
    print(f"  Word Pages: {summary['total_docx_pages']} (estimated)")
    print(f"  Total:      {summary['total_pages']}")
    
    # Print per-client summary table
    print("\n" + "-" * 80)
    print(f"{'Client Folder':<50} {'Docs':>8} {'Pages':>8}")
    print("-" * 80)
    
    for client_name, stats in sorted(clients.items()):
        # Truncate long names
        display_name = client_name[:47] + "..." if len(client_name) > 50 else client_name
        print(f"{display_name:<50} {stats['total_documents']:>8} {stats['total_pages']:>8}")
    
    print("-" * 80)
    print(f"{'TOTAL':<50} {summary['total_documents']:>8} {summary['total_pages']:>8}")
    print("=" * 80)


def save_detailed_report(results: dict, output_path: str):
    """Save detailed results to a JSON file."""
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(results, f, indent=2, ensure_ascii=False)
    print(f"\nDetailed report saved to: {output_path}")


def save_csv_report(results: dict, output_path: str):
    """Save summary to a CSV file."""
    with open(output_path, 'w', encoding='utf-8') as f:
        # Header
        f.write("Client Folder,PDF Count,DOCX Count,Total Documents,PDF Pages,DOCX Pages (est),Total Pages\n")
        
        for client_name, stats in sorted(results["clients"].items()):
            # Escape commas in client names
            safe_name = f'"{client_name}"' if ',' in client_name else client_name
            f.write(f"{safe_name},{stats['pdf_count']},{stats['docx_count']},{stats['total_documents']},"
                   f"{stats['pdf_pages']},{stats['docx_pages']},{stats['total_pages']}\n")
        
        # Total row
        summary = results["summary"]
        f.write(f"TOTAL,{summary['total_pdf_files']},{summary['total_docx_files']},{summary['total_documents']},"
               f"{summary['total_pdf_pages']},{summary['total_docx_pages']},{summary['total_pages']}\n")
    
    print(f"CSV report saved to: {output_path}")


if __name__ == "__main__":
    # Configuration
    TARGET_PATH = r"E:\f\Brdge AI-Projects\drive-download-20250903T145611Z-1-001"
    
    # Output paths
    OUTPUT_DIR = Path(__file__).parent / "output"
    OUTPUT_DIR.mkdir(exist_ok=True)
    
    JSON_OUTPUT = OUTPUT_DIR / "document_count_report.json"
    CSV_OUTPUT = OUTPUT_DIR / "document_count_report.csv"
    
    try:
        # Run analysis
        results = analyze_folder(TARGET_PATH)
        
        # Print summary
        print_summary(results)
        
        # Save reports
        save_detailed_report(results, str(JSON_OUTPUT))
        save_csv_report(results, str(CSV_OUTPUT))
        
        print("\nAnalysis complete!")
        
    except FileNotFoundError as e:
        print(f"Error: {e}")
    except Exception as e:
        print(f"An error occurred: {e}")
        raise
